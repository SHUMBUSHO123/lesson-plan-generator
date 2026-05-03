from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
doc.add_heading('Banner Examples for IsomoPlus', 0)

# 1. Top Banner Table
doc.add_heading('1. Top Banner (Ticker) – 10 Examples', level=1)
table1 = doc.add_table(rows=11, cols=10)
table1.style = 'Table Grid'
# headers
headers1 = ['#','Text','Icon','Badge','Badge color','Color','Bg','Scroll speed','Order','Show on']
for i, h in enumerate(headers1):
    table1.cell(0,i).text = h
# data (abbreviated for space; copy from the table above)
data1 = [
    ['1','🔥 Limited Offer: 40% off all premium plans – use code TEACH40','🔥','HOT','#e74c3c','#ffd700','#0d2b1f','44','0','All'],
    ['2','📚 New CBC resources added: Download free worksheets now!','📚','NEW','#27ae60','#fff','#1a3a5c','50','1','All'],
    ['3','⏳ Flash sale ends midnight: Get 3 months for the price of 2','⏳','SALE','#f39c12','#000','#f8f9fa','35','2','All'],
    ['4','🎓 Teacher appreciation week: 20% off any subscription','🎓','EVENT','#9b59b6','#fff','#2c3e50','60','3','All'],
    ['5','🏫 Back to school: Buy one term, get one free – limited slots','🏫','BOGO','#e67e22','#fff','#1e3c2c','55','4','All'],
    ['6','📱 New mobile app: Generate lessons on the go – try now','📱','NEW','#3498db','#fff','#0d1b2a','48','5','All'],
    ['7','⭐ Rated 4.9/5 by teachers: Join 10,000+ happy educators','⭐','TRUST','#f1c40f','#2c3e50','#ecf0f1','70','6','All'],
    ['8','💰 Earn money: Join our affiliate program – 30% recurring commission','💰','AFFILIATE','#e74c3c','#fff','#1a1a2e','52','7','All'],
    ['9','🏠 Property of the week: Perfect for school location – contact us','🏠','FEATURED','#16a085','#fff','#2c3e50','65','8','Landing'],
    ['10','📦 AliExpress deal: Up to 70% off teaching supplies – limited time','📦','DEAL','#e74c3c','#000','#fffae6','40','9','Index']
]
for i, row in enumerate(data1):
    for j, val in enumerate(row):
        table1.cell(i+1, j).text = val

doc.add_page_break()

# 2. Bottom Ad Table
doc.add_heading('2. Bottom Ad – 10 Examples', level=1)
table2 = doc.add_table(rows=11, cols=7)
table2.style = 'Table Grid'
headers2 = ['#','Title','Subtitle','Badge','Badge color','Url','Order','Show on']
for i, h in enumerate(headers2):
    table2.cell(0,i).text = h
data2 = [
    ['1','🚀 Upgrade to Premium','Unlimited lesson plans + PDF/DOCX export','SALE','#e74c3c','/pricing/','0','All'],
    ['2','📖 Free E‑book','"10 Secrets to Stress‑Free Lesson Planning"','FREE','#27ae60','/free-ebook/','1','Landing'],
    ['3','🛒 Shop Teacher Supplies','Up to 50% off on Amazon','AFFILIATE','#f39c12','https://amzn.to/example','2','All'],
    ['4','🏠 Rent a Classroom Space','Fully equipped for 30 students','FEATURED','#3498db','/rentals/','3','Index'],
    ['5','📱 Download Our App','Generate lessons from your phone','NEW','#9b59b6','/app','4','All'],
    ['6','💼 Sell Your Products','List on our marketplace – zero fees','HOT','#e67e22','/vendor-signup/','5','Pricing'],
    ['7','🎓 Grant Opportunity','Apply for free premium access','GRANT','#16a085','/grant/','6','Landing'],
    ['8','📦 AliExpress Choice','Best‑selling classroom gadgets','DEAL','#e74c3c','https://aliexpress.com/example','7','Index'],
    ['9','🏆 Win a Lifetime Plan','Enter our monthly giveaway','GIVEAWAY','#f1c40f','/giveaway/','8','All'],
    ['10','👩‍🏫 Become a Trainer','Join our certified trainer program','CAREER','#2c3e50','/trainer/','9','All']
]
for i, row in enumerate(data2):
    for j, val in enumerate(row):
        table2.cell(i+1, j).text = val

doc.add_page_break()

# 3. Hero Banner Table
doc.add_heading('3. Hero Banner (Carousel) – 10 Examples', level=1)
table3 = doc.add_table(rows=11, cols=12)
table3.style = 'Table Grid'
headers3 = ['#','Headline','Description','CTA text','CTA url','Badge','Badge color','Bg color','Bg color2','Duration','Order','Show on']
for i, h in enumerate(headers3):
    table3.cell(0,i).text = h
data3 = [
    ['1','🎯 Generate CBC Lesson Plans in 2 Minutes','AI‑powered, curriculum‑aligned, and ready to download','Start Free','/register/','NEW','#27ae60','#0d1b2a','#1a3a5c','8','0','All'],
    ['2','📉 Save 10+ Hours Every Week','Join thousands of teachers who use IsomoPlus','See Plans','/pricing/','SALE','#e74c3c','#1a1a2e','#16213e','7','1','All'],
    ['3','🏫 Special Offer for Schools','Get 50% off when you sign up 5+ teachers','Contact Us','/school-plan/','SCHOOL','#3498db','#0d3a1a','#2a6b1a','9','2','Landing'],
    ['4','🎓 100% Free – No Credit Card Required','Start with 3 free lesson plans today','Register Now','/register/','FREE','#16a085','#2c3e50','#34495e','6','3','Index'],
    ['5','📚 Affiliate Program – Earn 30% Recurring','Promote IsomoPlus and get paid monthly','Join Now','/affiliate/','AFFILIATE','#f39c12','#1e2a3a','#0d1b2a','10','4','Pricing'],
    ['6','🏠 Property Ad: Ideal for Learning Centers','Spacious, quiet, near schools – lease today','View Property','/property/','FEATURED','#e67e22','#5d3a1a','#8b5a2b','12','5','Landing'],
    ['7','📦 AliExpress Top Picks','Best selling classroom tools – up to 70% off','Shop Now','https://aliexpress.com/example','DEAL','#e74c3c','#1a1a2e','#2c3e50','8','6','Index'],
    ['8','🎮 Gamified Learning for Students','Engage your class with interactive quizzes','Learn More','/gamification/','NEW','#9b59b6','#2c0b3a','#4a1a6b','7','7','All'],
    ['9','💼 Teacher Side Hustle','Sell your own lesson plans on our marketplace','Become a Seller','/seller/','EARN','#16a085','#0d3a3a','#1a6b6b','9','8','Pricing'],
    ['10','🎉 Summer Special: 50% Off Annual Plan','Limited time – upgrade now and save big','Claim Offer','/pricing/?promo=summer50','HOT','#e74c3c','#6b1a1a','#a32a2a','5','9','All']
]
for i, row in enumerate(data3):
    for j, val in enumerate(row):
        table3.cell(i+1, j).text = val

doc.save('IsomoPlus_Banner_Examples.docx')
print("DOCX file created: IsomoPlus_Banner_Examples.docx")