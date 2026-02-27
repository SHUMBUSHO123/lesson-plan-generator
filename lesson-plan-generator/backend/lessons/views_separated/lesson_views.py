print("🔥 lesson_views.py START LOADING")

from django.shortcuts import get_object_or_404
from django.http import JsonResponse, HttpResponse
from django.utils import timezone
from django.template.loader import render_to_string
from datetime import timedelta
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import AllowAny
from rest_framework.response import Response
from rest_framework import status
from lessons.helpers import build_lesson_plan_context, prepare_lesson_plan_context
from lessons.models import (
    UserProfile, Lesson, Unit, FREE_LESSON_LIMIT, SUBSCRIPTION_LIMITS,
    TeachingStrategy, LessonStrategyStep, GeneratedLessonPlan
)
from lessons.utils.resolve_profile import resolve_profile
from lessons.services.lesson_engine import build_strategy_steps
import random
import io
import zipfile
from django.db.models import F

# ── ReportLab — pure Python PDF, zero system dependencies ──────────────────
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import mm
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle,
    Paragraph, Spacer, KeepTogether
)

# ── python-docx — pure Python Word, zero system dependencies ───────────────
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────────────────────────────────────
# Shared PDF colours & helpers
# ─────────────────────────────────────────────────────────────────────────────
HEAD_BG  = colors.HexColor('#d9e1f2')
LABEL_BG = colors.HexColor('#f2f2f2')
DESC_BG  = colors.HexColor('#f9f9f9')
BORDER   = colors.HexColor('#333333')


def _style(font='Helvetica', size=7.5, leading=9.5, align=TA_LEFT, italic=False):
    return ParagraphStyle(
        name=f'{font}-{size}-{align}',
        fontName='Helvetica-Oblique' if italic else font,
        fontSize=size,
        leading=leading,
        alignment=align,
    )


def _base_table_style():
    return TableStyle([
        ('GRID',          (0, 0), (-1, -1), 0.5,  BORDER),
        ('VALIGN',        (0, 0), (-1, -1), 'TOP'),
        ('FONTNAME',      (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE',      (0, 0), (-1, -1), 7.5),
        ('LEADING',       (0, 0), (-1, -1), 9.5),
        ('LEFTPADDING',   (0, 0), (-1, -1), 3),
        ('RIGHTPADDING',  (0, 0), (-1, -1), 3),
        ('TOPPADDING',    (0, 0), (-1, -1), 2),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
    ])


# ─────────────────────────────────────────────────────────────────────────────
# Shared Word (docx) helpers
# ─────────────────────────────────────────────────────────────────────────────
def _hex_to_rgb(hex_str):
    """Convert '#d9e1f2' to RGBColor(217, 225, 242)."""
    h = hex_str.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_cell_bg(cell, hex_color):
    """Set a Word table cell background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color.lstrip('#'))
    tcPr.append(shd)


def _set_cell_borders(cell, color='333333'):
    """Add borders to a single Word table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _bold_run(para, text, size_pt=8):
    run = para.add_run(text)
    run.bold      = True
    run.font.size = Pt(size_pt)
    return run


def _normal_run(para, text, size_pt=8, italic=False):
    run = para.add_run(text)
    run.bold        = False
    run.italic      = italic
    run.font.size   = Pt(size_pt)
    return run


def _info_row(table, label, value, label_bg='f2f2f2', head_bg='d9e1f2'):
    """Add a 2-column label|value row to the Word info table."""
    row        = table.add_row()
    label_cell = row.cells[0]
    value_cell = row.cells[1]

    label_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    value_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    _set_cell_bg(label_cell, label_bg)
    _set_cell_borders(label_cell)
    _set_cell_borders(value_cell)

    lp = label_cell.paragraphs[0]
    _bold_run(lp, label, 8)

    vp = value_cell.paragraphs[0]
    _normal_run(vp, str(value) if value else '', 8)

    return row


# ─────────────────────────────────────────────────────────────────────────────
def use_lesson_atomic(profile):
    profile.lessons_used = F('lessons_used') + 1
    profile.save(update_fields=['lessons_used'])
    profile.refresh_from_db()


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────
def get_or_create_guest_profile(request):
    profile = resolve_profile(request)
    if profile.is_guest and profile.lesson_limit is None:
        profile.lesson_limit = FREE_LESSON_LIMIT
        profile.save(update_fields=['lesson_limit'])
    return profile


# ─────────────────────────────────────────────────────────────────────────────
# Lesson Generation
# ─────────────────────────────────────────────────────────────────────────────
@api_view(['POST'])
@permission_classes([AllowAny])
def generate_lesson_plan(request):
    profile = resolve_profile(request)
    profile.expire_subscription_if_needed()

    if profile.is_premium and not profile.is_active:
        return Response({"error": "Premium account inactive."}, status=403)

    if not profile.can_generate_plan():
        return Response({"redirect": "/pricing/"}, status=403)

    print("REQUEST.DATA:", request.data)

    try:
        context = build_lesson_plan_context(request.data, profile)
    except PermissionError as e:
        return Response({"error": str(e), "redirect": "/pricing/"}, status=403)
    except Exception as e:
        import traceback; traceback.print_exc()
        return Response({"error": f"Failed to build context: {str(e)}"}, status=500)

    try:
        use_lesson_atomic(profile)
        context['lesson_info']['lessons_used'] = profile.lessons_used
    except Exception as e:
        import traceback; traceback.print_exc()
        return Response({"error": f"Failed to record lesson usage: {str(e)}"}, status=500)

    try:
        html_content = render_to_string("partials/lesson_plan_table.html", context)
    except Exception as e:
        import traceback; traceback.print_exc()
        return Response({"error": f"Failed to render lesson plan: {str(e)}"}, status=500)

    if request.user.is_authenticated:
        try:
            lesson_info = context.get('lesson_info', {})
            GeneratedLessonPlan.objects.create(
                profile       = profile,
                subject       = lesson_info.get('subject', request.data.get('subject', '')),
                unit_title    = lesson_info.get('unit_title', ''),
                lesson_title  = lesson_info.get('lesson_title', '')[:300],
                class_name    = lesson_info.get('class_name', request.data.get('class', '')),
                term          = request.data.get('term', ''),
                html_snapshot = html_content,
            )
        except Exception as save_err:
            print(f"⚠️  Could not save GeneratedLessonPlan: {save_err}")

    return Response({"html": html_content})


# ─────────────────────────────────────────────────────────────────────────────
# Access Check
# ─────────────────────────────────────────────────────────────────────────────
@api_view(['GET'])
@permission_classes([AllowAny])
def check_access(request):
    profile = get_or_create_guest_profile(request)
    if profile.subscription_expiry:
        profile.expire_subscription_if_needed()

    is_premium   = profile.is_premium
    can_generate = profile.is_active and (is_premium or profile.can_generate_plan())

    return Response({
        "can_generate": can_generate,
        "is_premium":   is_premium,
        "lessons_used": profile.lessons_used,
        "lesson_limit": profile.get_current_lesson_limit(),
        "is_active":    profile.is_active,
        "message":      "Access check complete",
    })


# ─────────────────────────────────────────────────────────────────────────────
# Subscription Check
# ─────────────────────────────────────────────────────────────────────────────
@api_view(['GET'])
def check_subscription(request):
    device_id = request.GET.get('device_id')
    if not device_id:
        return JsonResponse({"active": False, "expires_at": None})
    profile = UserProfile.objects.filter(device_id=device_id).first()
    if not profile:
        return JsonResponse({"active": False, "expires_at": None})
    if profile.subscription_expiry:
        profile.expire_subscription_if_needed()
    return JsonResponse({"active": profile.is_premium, "expires_at": profile.subscription_expiry})


# ─────────────────────────────────────────────────────────────────────────────
# User Dashboard
# ─────────────────────────────────────────────────────────────────────────────
@api_view(['GET'])
@permission_classes([AllowAny])
def get_user_dashboard(request):
    is_authenticated = request.user.is_authenticated

    if is_authenticated:
        profile, _ = UserProfile.objects.get_or_create(user=request.user)
        profile.expire_subscription_if_needed()

        limit = profile.get_current_lesson_limit()
        used  = profile.lessons_used
        remaining = None if (profile.lesson_limit is None and profile.is_premium) \
                         else max(0, limit - used)

        payload = {
            "is_authenticated":    True,
            "is_premium":          profile.is_premium,
            "subscription_plan":   profile.subscription_plan,
            "subscription_expiry": (
                profile.subscription_expiry.isoformat()
                if profile.subscription_expiry else None
            ),
            "lessons_used":      used,
            "lesson_limit":      limit,
            "remaining":         remaining,
            "can_download_pdf":  profile.can_download_pdf,
            "can_download_docx": profile.can_download_docx,
            "can_copy_word":     profile.can_copy_word,
            "recent_plans":      [],
        }

        plans = GeneratedLessonPlan.objects.filter(
            profile=profile
        ).order_by('-created_at')[:10]

        payload["recent_plans"] = [
            {
                "id":           p.id,
                "subject":      p.subject,
                "unit_title":   p.unit_title,
                "lesson_title": p.lesson_title,
                "class_name":   p.class_name,
                "term":         p.term,
                "created_at":   p.created_at.strftime("%d %b %Y, %H:%M"),
            }
            for p in plans
        ]

    else:
        profile   = resolve_profile(request)
        limit     = profile.get_current_lesson_limit()
        used      = profile.lessons_used
        remaining = max(0, limit - used)

        payload = {
            "is_authenticated":    False,
            "is_premium":          False,
            "subscription_plan":   None,
            "subscription_expiry": None,
            "lessons_used":        used,
            "lesson_limit":        limit,
            "remaining":           remaining,
            "can_download_pdf":    False,
            "can_download_docx":   False,
            "can_copy_word":       False,
            "recent_plans":        [],
        }

    return Response(payload)


# ─────────────────────────────────────────────────────────────────────────────
# Bulk ZIP Download
# ─────────────────────────────────────────────────────────────────────────────
@api_view(['POST'])
def bulk_zip_download(request):
    if not request.user.is_authenticated:
        return Response({"error": "Login required."}, status=401)

    profile, _ = UserProfile.objects.get_or_create(user=request.user)
    plan_ids   = request.data.get('plan_ids', [])
    if not plan_ids or not isinstance(plan_ids, list):
        return Response({"error": "Provide a list of plan_ids."}, status=400)

    plans = GeneratedLessonPlan.objects.filter(id__in=plan_ids, profile=profile)
    if not plans.exists():
        return Response({"error": "No matching plans found."}, status=404)

    PLAN_CSS = """
        @page { size: A4 portrait; margin: 8mm; }
        *, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }
        html, body { width: 100%; background: #fff; color: #000; }
        body { font-family: Arial, sans-serif; font-size: 7.5pt; line-height: 1.15; }
        table { border-collapse: collapse; width: 100%; margin-bottom: 4px;
                page-break-inside: avoid; break-inside: avoid; }
        td, th { border: 1px solid #333; padding: 2px 4px; vertical-align: top; }
        th { background-color: #d9e1f2; font-weight: bold; text-align: center; }
        tr { page-break-inside: avoid; break-inside: avoid; }
        .label { font-weight: bold; background-color: #f2f2f2; white-space: nowrap; }
        .competence-integration { font-style: italic; font-size: 7.5pt; color: #222; margin-top: 2px; }
        .eval-options { list-style: none; padding: 0; margin: 0; }
        .eval-options li { margin-bottom: 2px; font-size: 7.5pt; }
        .eval-options li::before { content: "\\2610\\00a0"; font-size: 9pt; }
        .lp-wrap, .lp-wrap * { page-break-inside: avoid; break-inside: avoid; }
    """

    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        for plan in plans:
            full_html = f"""<!DOCTYPE html>
<html lang="en"><head><meta charset="UTF-8">
<title>{plan.subject} - {plan.lesson_title}</title>
<style>{PLAN_CSS}</style></head>
<body>{plan.html_snapshot}</body></html>"""
            safe_subject = "".join(c for c in plan.subject     if c.isalnum() or c in " _-")[:30].strip()
            safe_lesson  = "".join(c for c in plan.lesson_title if c.isalnum() or c in " _-")[:40].strip()
            date_str     = plan.created_at.strftime("%Y%m%d")
            filename     = f"{date_str}_{safe_subject}_{safe_lesson}.html".replace(" ", "_")
            zf.writestr(filename, full_html)

    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/zip')
    response['Content-Disposition'] = 'attachment; filename="CBC_Lesson_Plans.zip"'
    return response


# ─────────────────────────────────────────────────────────────────────────────
# Prefill API
# ─────────────────────────────────────────────────────────────────────────────
@api_view(['GET'])
def get_user_prefill(request):
    profile      = get_or_create_guest_profile(request)
    unit_id      = request.GET.get("unit_id")
    unit_display = ""
    lesson_title = ""

    if unit_id:
        unit = Unit.objects.filter(id=unit_id, is_active=True).first()
        if unit:
            unit_display = f"Unit {unit.number} - {unit.title}"
            first_lesson = unit.lessons.filter(is_active=True).order_by("number").first()
            if first_lesson:
                lesson_title = first_lesson.title

    return Response({
        "schoolName":  profile.school_name or "",
        "teacherName": (
            request.user.get_full_name()
            if request.user.is_authenticated
            else profile.teacher_name or ""
        ),
        "term":        profile.default_term or "I",
        "classSize":   profile.class_size or "",
        "references":  profile.references or "",
        "unitTitle":   unit_display,
        "lessonTitle": lesson_title,
    })


# ─────────────────────────────────────────────────────────────────────────────
# PDF Download — ReportLab
# Pure Python, zero system dependencies, works on Windows & Linux identically
# ─────────────────────────────────────────────────────────────────────────────
@api_view(['POST'])
@permission_classes([AllowAny])
def download_lesson_pdf(request):
    """
    Professional single-page A4 PDF via ReportLab.
    Pure Python — pip install reportlab is ALL you need.
    No GTK3, no apt-get, no system libraries.
    """
    profile = resolve_profile(request)
    profile.expire_subscription_if_needed()

    if not profile.is_premium:
        return Response(
            {"error": "PDF download requires a premium subscription.",
             "redirect": "/pricing/"},
            status=403
        )

    try:
        context = build_lesson_plan_context(request.data, profile)
    except PermissionError as e:
        return Response({"error": str(e), "redirect": "/pricing/"}, status=403)
    except Exception as e:
        return Response({"error": f"Failed to build context: {str(e)}"}, status=500)

    lesson_info             = context.get('lesson_info', {})
    unit_info               = context.get('unit_info', {})
    steps                   = context.get('steps', [])
    eval_options            = context.get('self_evaluation_options', [])
    lesson_description      = context.get('lesson_description', '')
    instructional_objective = context.get('instructional_objective', '')
    special_needs           = context.get('special_needs', '')

    buffer = io.BytesIO()
    doc    = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        topMargin=10*mm, bottomMargin=10*mm,
        leftMargin=10*mm, rightMargin=10*mm,
    )
    page_w = A4[0] - 20*mm

    normal   = _style('Helvetica',      7.5, 9.5)
    bold_s   = _style('Helvetica-Bold', 7.5, 9.5)
    title_s  = _style('Helvetica-Bold', 11,  14, TA_CENTER)
    small    = _style('Helvetica',      7.0, 9.0)
    italic_s = _style('Helvetica',      7.0, 9.0, italic=True)

    story = []
    story.append(Paragraph('LESSON PLAN', title_s))
    story.append(Spacer(1, 3*mm))

    # School & teacher row
    school_tbl = Table(
        [[
            Paragraph(f"<b>School Name:</b> {lesson_info.get('school_name','')}", normal),
            Paragraph(f"<b>Teacher's Name:</b> {lesson_info.get('teacher_name','')}", normal),
        ]],
        colWidths=[page_w * 0.5, page_w * 0.5]
    )
    school_tbl.setStyle(TableStyle([
        ('VALIGN',      (0,0), (-1,-1), 'TOP'),
        ('LEFTPADDING', (0,0), (-1,-1), 0),
    ]))
    story.append(school_tbl)
    story.append(Spacer(1, 2*mm))

    # Header grid
    cw8 = page_w / 8
    hdr_tbl = Table(
        [
            [Paragraph(f'<b>{h}</b>', bold_s) for h in
             ['Term','Date','Subject','Class','Unit No','Lesson No','Duration','Class size']],
            [
                Paragraph(str(lesson_info.get('term',        '')), normal),
                Paragraph(str(lesson_info.get('date',        '')), normal),
                Paragraph(str(lesson_info.get('subject',     '')), normal),
                Paragraph(str(lesson_info.get('class_name',  '')), normal),
                Paragraph(f"Unit {unit_info.get('unit_number','')}", normal),
                Paragraph(f"{context.get('lesson_number','')} of {unit_info.get('total_lessons','')}", normal),
                Paragraph(f"{context.get('lesson_duration','')} min", normal),
                Paragraph(str(lesson_info.get('class_size',  '')), normal),
            ],
        ],
        colWidths=[cw8] * 8
    )
    hs = _base_table_style()
    hs.add('BACKGROUND', (0,0), (-1,0), HEAD_BG)
    hs.add('FONTNAME',   (0,0), (-1,0), 'Helvetica-Bold')
    hs.add('ALIGN',      (0,0), (-1,0), 'CENTER')
    hdr_tbl.setStyle(hs)
    story.append(hdr_tbl)

    # Info table
    def irow(label, value):
        return [Paragraph(label, bold_s), Paragraph(str(value) if value else '', normal)]

    info_tbl = Table(
        [
            irow('Special Educational Needs', special_needs),
            irow('Unit Title',                unit_info.get('unit_title', '')),
            irow('Key Unit Competence',        unit_info.get('key_unit_competence', '')),
            irow('Title of the Lesson',        lesson_info.get('lesson_title', '')),
            irow('Instructional Objective',    instructional_objective),
            irow('Plan for this Class',        lesson_info.get('location_plan', '')),
            irow('Learning Materials',         lesson_info.get('materials', '')),
            irow('References',                 lesson_info.get('references', '')),
        ],
        colWidths=[page_w * 0.28, page_w * 0.72]
    )
    info_s = _base_table_style()
    info_s.add('BACKGROUND', (0, 0), (0, -1), LABEL_BG)
    info_tbl.setStyle(info_s)
    story.append(info_tbl)

    # Activities table
    cw = [page_w*0.13, page_w*0.28, page_w*0.28, page_w*0.31]
    act_data = [
        [
            Paragraph('<b>Timing / Step</b>', bold_s),
            Paragraph('<b>Description of Teaching and Learning Activity</b>', bold_s),
            '',
            Paragraph('<b>Generic Competences and How They Are Integrated</b>', bold_s),
        ],
        [
            '',
            Paragraph(f'<i><b>Lesson Description:</b> {lesson_description}</i>', italic_s),
            '', '',
        ],
        [
            '',
            Paragraph('<b>Teacher Activities</b>', bold_s),
            Paragraph('<b>Learner Activities</b>', bold_s),
            '',
        ],
    ]

    for step in steps:
        comp = f"<b>{step.get('competence', '')}</b>"
        if step.get('competence_integration'):
            comp += f"<br/><i><b>Integration:</b> {step['competence_integration']}</i>"
        act_data.append([
            Paragraph(f"<b>{step.get('title','')}</b><br/>{step.get('duration_minutes','')} min", normal),
            Paragraph(step.get('teacher_activity', ''), normal),
            Paragraph(step.get('learner_activity',  ''), normal),
            Paragraph(comp, small),
        ])

    eval_text = '<br/>'.join(f'&#9744;  {opt}' for opt in eval_options)
    last_row  = len(act_data)
    act_data.append([
        Paragraph('<b>Teacher Self-Evaluation</b>', bold_s),
        Paragraph(eval_text, small),
        '', '',
    ])

    act_tbl = Table(act_data, colWidths=cw)
    act_s   = _base_table_style()
    act_s.add('BACKGROUND', (0,0), (-1,0), HEAD_BG)
    act_s.add('FONTNAME',   (0,0), (-1,0), 'Helvetica-Bold')
    act_s.add('ALIGN',      (0,0), (-1,0), 'CENTER')
    act_s.add('VALIGN',     (0,0), (-1,0), 'MIDDLE')
    act_s.add('SPAN',       (1,0), (2,0))
    act_s.add('BACKGROUND', (1,1), (2,1), DESC_BG)
    act_s.add('SPAN',       (1,1), (2,1))
    act_s.add('SPAN',       (0,0), (0,2))
    act_s.add('ALIGN',      (0,0), (0,2), 'CENTER')
    act_s.add('VALIGN',     (0,0), (0,2), 'MIDDLE')
    act_s.add('SPAN',       (3,0), (3,2))
    act_s.add('ALIGN',      (3,0), (3,2), 'CENTER')
    act_s.add('VALIGN',     (3,0), (3,2), 'MIDDLE')
    act_s.add('BACKGROUND', (1,2), (2,2), HEAD_BG)
    act_s.add('FONTNAME',   (1,2), (2,2), 'Helvetica-Bold')
    act_s.add('ALIGN',      (1,2), (2,2), 'CENTER')
    act_s.add('BACKGROUND', (0, last_row), (0, last_row), LABEL_BG)
    act_s.add('SPAN',       (1, last_row), (3, last_row))
    act_tbl.setStyle(act_s)
    story.append(KeepTogether([act_tbl]))

    try:
        doc.build(story)
    except Exception as e:
        return Response({"error": f"PDF build failed: {str(e)}"}, status=500)

    buffer.seek(0)
    safe_subject = "".join(c for c in lesson_info.get('subject', 'Lesson') if c.isalnum() or c in " _-")[:20].strip()
    safe_title   = "".join(c for c in lesson_info.get('lesson_title', '')  if c.isalnum() or c in " _-")[:30].strip()
    filename     = f"CBC_{safe_subject}_{safe_title}.pdf".replace(" ", "_")

    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


# ─────────────────────────────────────────────────────────────────────────────
# Word Download — python-docx
# Pure Python, zero system dependencies
# Produces a real .docx — fully editable in Word, LibreOffice, Google Docs
# pip install python-docx
# ─────────────────────────────────────────────────────────────────────────────
@api_view(['POST'])
@permission_classes([AllowAny])
def download_lesson_word(request):
    """
    Professional editable .docx via python-docx.
    Pure Python — pip install python-docx is ALL you need.
    No GTK3, no apt-get, no system libraries.
    Produces real Word tables with bold labels, coloured headers,
    and fully editable text — not a screenshot, not HTML.
    """
    profile = resolve_profile(request)
    profile.expire_subscription_if_needed()

    if not profile.is_premium:
        return Response(
            {"error": "Word download requires a premium subscription.",
             "redirect": "/pricing/"},
            status=403
        )

    try:
        context = build_lesson_plan_context(request.data, profile)
    except PermissionError as e:
        return Response({"error": str(e), "redirect": "/pricing/"}, status=403)
    except Exception as e:
        return Response({"error": f"Failed to build context: {str(e)}"}, status=500)

    lesson_info             = context.get('lesson_info', {})
    unit_info               = context.get('unit_info', {})
    steps                   = context.get('steps', [])
    eval_options            = context.get('self_evaluation_options', [])
    lesson_description      = context.get('lesson_description', '')
    instructional_objective = context.get('instructional_objective', '')
    special_needs           = context.get('special_needs', '')

    # ── Build Word document ───────────────────────────────────────────────
    doc = Document()

    # Page size A4, tight margins
    section = doc.sections[0]
    section.page_width    = Cm(21.0)
    section.page_height   = Cm(29.7)
    section.top_margin    = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin   = Cm(1.0)
    section.right_margin  = Cm(1.0)

    # ── TITLE ─────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run  = title_para.add_run('LESSON PLAN')
    title_run.bold      = True
    title_run.font.size = Pt(13)

    # ── SCHOOL & TEACHER ──────────────────────────────────────────────────
    school_tbl = doc.add_table(rows=1, cols=2)
    school_tbl.style = 'Table Grid'
    school_tbl.autofit = True

    left  = school_tbl.cell(0, 0)
    right = school_tbl.cell(0, 1)
    _set_cell_borders(left)
    _set_cell_borders(right)

    lp = left.paragraphs[0]
    _bold_run(lp, 'School Name: ', 8)
    _normal_run(lp, lesson_info.get('school_name', ''), 8)

    rp = right.paragraphs[0]
    _bold_run(rp, "Teacher's Name: ", 8)
    _normal_run(rp, lesson_info.get('teacher_name', ''), 8)

    doc.add_paragraph()   # small spacer

    # ── HEADER GRID ───────────────────────────────────────────────────────
    # Term | Date | Subject | Class | Unit No | Lesson No | Duration | Size
    hdr_tbl = doc.add_table(rows=2, cols=8)
    hdr_tbl.style  = 'Table Grid'
    hdr_tbl.autofit = True

    headers = ['Term','Date','Subject','Class','Unit No','Lesson No','Duration','Class size']
    values  = [
        lesson_info.get('term',       ''),
        lesson_info.get('date',       ''),
        lesson_info.get('subject',    ''),
        lesson_info.get('class_name', ''),
        f"Unit {unit_info.get('unit_number','')}",
        f"{context.get('lesson_number','')} of {unit_info.get('total_lessons','')}",
        f"{context.get('lesson_duration','')} min",
        lesson_info.get('class_size', ''),
    ]

    for i, h in enumerate(headers):
        cell = hdr_tbl.cell(0, i)
        _set_cell_bg(cell, 'd9e1f2')
        _set_cell_borders(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _bold_run(p, h, 7.5)

    for i, v in enumerate(values):
        cell = hdr_tbl.cell(1, i)
        _set_cell_borders(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _normal_run(p, str(v), 7.5)

    doc.add_paragraph()

    # ── INFO TABLE ────────────────────────────────────────────────────────
    info_tbl = doc.add_table(rows=0, cols=2)
    info_tbl.style  = 'Table Grid'
    info_tbl.autofit = True

    # Set column widths: 28% label, 72% value
    for row_data in [
        ('Special Educational Needs', special_needs),
        ('Unit Title',                unit_info.get('unit_title', '')),
        ('Key Unit Competence',       unit_info.get('key_unit_competence', '')),
        ('Title of the Lesson',       lesson_info.get('lesson_title', '')),
        ('Instructional Objective',   instructional_objective),
        ('Plan for this Class',       lesson_info.get('location_plan', '')),
        ('Learning Materials',        lesson_info.get('materials', '')),
        ('References',                lesson_info.get('references', '')),
    ]:
        _info_row(info_tbl, row_data[0], row_data[1])

    doc.add_paragraph()

    # ── ACTIVITIES TABLE ──────────────────────────────────────────────────
    # Rows: header(3) + steps + self-eval
    total_rows = 3 + len(steps) + 1
    act_tbl    = doc.add_table(rows=total_rows, cols=4)
    act_tbl.style  = 'Table Grid'
    act_tbl.autofit = True

    # Row 0 — main column headers
    col_headers = [
        'Timing / Step',
        'Description of Teaching and Learning Activity',
        '',
        'Generic Competences and How They Are Integrated',
    ]
    for i, h in enumerate(col_headers):
        cell = act_tbl.cell(0, i)
        _set_cell_bg(cell, 'd9e1f2')
        _set_cell_borders(cell)
        if h:
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _bold_run(p, h, 7.5)

    # Merge teacher+learner header columns (cols 1+2 in row 0)
    act_tbl.cell(0, 1).merge(act_tbl.cell(0, 2))

    # Row 1 — lesson description (merge cols 1+2)
    desc_cell = act_tbl.cell(1, 1)
    desc_cell.merge(act_tbl.cell(1, 2))
    _set_cell_bg(desc_cell, 'f9f9f9')
    _set_cell_borders(desc_cell)
    dp = desc_cell.paragraphs[0]
    _bold_run(dp,   'Lesson Description: ', 7.5)
    _normal_run(dp, lesson_description, 7.5, italic=True)

    # Timing/Step cell spans rows 0-2 (merge)
    act_tbl.cell(0, 0).merge(act_tbl.cell(2, 0))
    # Competences cell spans rows 0-2
    act_tbl.cell(0, 3).merge(act_tbl.cell(2, 3))

    # Row 2 — sub-column headers
    for i, h in enumerate(['', 'Teacher Activities', 'Learner Activities', '']):
        cell = act_tbl.cell(2, i)
        _set_cell_bg(cell, 'd9e1f2')
        _set_cell_borders(cell)
        if h:
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _bold_run(p, h, 7.5)

    # Step rows (start at row 3)
    for idx, step in enumerate(steps):
        row_idx = 3 + idx

        # Timing cell
        timing_cell = act_tbl.cell(row_idx, 0)
        _set_cell_borders(timing_cell)
        tp = timing_cell.paragraphs[0]
        _bold_run(tp,   f"{step.get('title','')}\n", 7.5)
        _normal_run(tp, f"{step.get('duration_minutes','')} min", 7.5)

        # Teacher cell
        teacher_cell = act_tbl.cell(row_idx, 1)
        _set_cell_borders(teacher_cell)
        _normal_run(teacher_cell.paragraphs[0], step.get('teacher_activity', ''), 7.5)

        # Learner cell
        learner_cell = act_tbl.cell(row_idx, 2)
        _set_cell_borders(learner_cell)
        _normal_run(learner_cell.paragraphs[0], step.get('learner_activity', ''), 7.5)

        # Competence cell
        comp_cell = act_tbl.cell(row_idx, 3)
        _set_cell_borders(comp_cell)
        cp = comp_cell.paragraphs[0]
        _bold_run(cp, step.get('competence', ''), 7.5)
        if step.get('competence_integration'):
            cp.add_run('\n')
            _bold_run(cp,   'Integration: ', 7)
            _normal_run(cp, step['competence_integration'], 7, italic=True)

    # Self-evaluation row
    eval_row_idx  = 3 + len(steps)
    eval_label    = act_tbl.cell(eval_row_idx, 0)
    eval_val_cell = act_tbl.cell(eval_row_idx, 1)

    _set_cell_bg(eval_label, 'f2f2f2')
    _set_cell_borders(eval_label)
    ep = eval_label.paragraphs[0]
    _bold_run(ep, 'Teacher Self-Evaluation', 7.5)

    # Merge value across cols 1-3
    eval_val_cell.merge(act_tbl.cell(eval_row_idx, 3))
    _set_cell_borders(eval_val_cell)
    evp = eval_val_cell.paragraphs[0]
    for opt in eval_options:
        evp.add_run(f'☐  {opt}\n').font.size = Pt(7.5)

    # ── Stream response ───────────────────────────────────────────────────
    buffer = io.BytesIO()
    try:
        doc.save(buffer)
    except Exception as e:
        return Response({"error": f"Word build failed: {str(e)}"}, status=500)

    buffer.seek(0)
    safe_subject = "".join(c for c in lesson_info.get('subject', 'Lesson') if c.isalnum() or c in " _-")[:20].strip()
    safe_title   = "".join(c for c in lesson_info.get('lesson_title', '')  if c.isalnum() or c in " _-")[:30].strip()
    filename     = f"CBC_{safe_subject}_{safe_title}.docx".replace(" ", "_")

    response = HttpResponse(
        buffer,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response